"""Paleta de cores e helpers de estilo corporativo BScash."""
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta ────────────────────────────────────────────────────────────────────
LARANJA   = RGBColor(0xFF, 0x6F, 0x00)
PRETO     = RGBColor(0x1A, 0x1A, 0x1A)
CINZA     = RGBColor(0x55, 0x55, 0x55)
BRANCO    = RGBColor(0xFF, 0xFF, 0xFF)
CINZA_BG  = RGBColor(0xF5, 0xF5, 0xF5)
AZUL_HDR  = RGBColor(0x1F, 0x39, 0x64)   # cabeçalho/rodapé

# Complexidade
COR_ALTA  = RGBColor(0xFF, 0xEB, 0xEE)
COR_MEDIA = RGBColor(0xFF, 0xF3, 0xE0)
COR_BAIXA = RGBColor(0xE8, 0xF5, 0xE9)
COR_CRITICA = RGBColor(0xFF, 0xCC, 0xBC)


def comp_color(val: str) -> RGBColor:
    v = (val or "").upper()
    if "ALTA" in v or "ALTO" in v or "CRÍTICA" in v or "CRITICA" in v:
        return COR_ALTA
    if "MÉDIA" in v or "MEDIA" in v or "MÉDIO" in v or "MEDIO" in v:
        return COR_MEDIA
    return COR_BAIXA


def shade_cell(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    tcPr.append(shd)


def cell_text(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    if align:
        p.alignment = align
    run = p.add_run(str(text or "—"))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color or CINZA


def heading(doc, text, level=1, numbered=None):
    label = f"{numbered}. {text}" if numbered else text
    p = doc.add_heading(label, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = LARANJA if level == 1 else PRETO
        run.font.size = Pt(15 if level == 1 else 12 if level == 2 else 11)
    return p


def para(doc, text, bold=False, size=10, color=None, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    run = p.add_run(str(text or ""))
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color or CINZA
    return p


def table(doc, headers, rows, col_widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade_cell(c, PRETO)
        cell_text(c, h, bold=True, size=9, color=BRANCO)
    for ri, row in enumerate(rows):
        tr = t.add_row()
        bg = CINZA_BG if ri % 2 == 0 else BRANCO
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            if isinstance(val, tuple):
                shade_cell(c, val[1])
                cell_text(c, val[0], size=font_size)
            else:
                shade_cell(c, bg)
                cell_text(c, val, size=font_size)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def page_break(doc):
    doc.add_page_break()


def set_margins(doc, top=2, bottom=2, left=2.5, right=2.5):
    for s in doc.sections:
        s.top_margin = Cm(top)
        s.bottom_margin = Cm(bottom)
        s.left_margin = Cm(left)
        s.right_margin = Cm(right)
