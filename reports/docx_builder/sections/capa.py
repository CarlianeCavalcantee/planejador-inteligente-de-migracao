from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..styles import LARANJA, PRETO, CINZA, page_break


def build(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(8)

    r = p.add_run(f"{data.get('sistema_escopo', 'BScash')}\n\n")
    r.font.size = Pt(18); r.bold = True; r.font.color.rgb = LARANJA; r.font.name = "Calibri"

    r2 = p.add_run("DOCUMENTO DE ARQUITETURA (SPEC)\n")
    r2.font.size = Pt(22); r2.bold = True; r2.font.color.rgb = PRETO; r2.font.name = "Calibri"

    r3 = p.add_run("ADEQUAÇÃO AO CNPJ ALFANUMÉRICO\n")
    r3.font.size = Pt(15); r3.bold = True; r3.font.color.rgb = LARANJA; r3.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)

    ver = data.get("spec_versao", "1.0")
    dt = data.get("data_execucao", data.get("data_geracao", "—"))[:10] if data.get("data_execucao") else "—"
    scan_id = data.get("scan_id", "")
    limite = data.get("data_limite_migracao", "—")

    scan_id_str = f"  |  ID: {scan_id}" if scan_id else ""
    r4 = sub.add_run(f"Versão: {ver}  |  Gerado em: {dt}{scan_id_str}  |  Prazo limite: {limite}\n")
    r4.font.size = Pt(10); r4.font.color.rgb = CINZA; r4.font.name = "Calibri"

    r5 = sub.add_run("Autor: Engenharia de Software  |  Status: Em Elaboração")
    r5.font.size = Pt(10); r5.italic = True; r5.font.color.rgb = CINZA; r5.font.name = "Calibri"

    page_break(doc)
